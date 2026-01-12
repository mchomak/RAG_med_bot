"""
RAG Engine для работы с медицинской документацией
Основной класс для индексации документов и генерации ответов
"""

import os
import sys
from pathlib import Path
from typing import List, Optional, Dict, Any
import warnings
warnings.filterwarnings('ignore')

# LlamaIndex imports
from llama_index.core import (
    VectorStoreIndex,
    SimpleDirectoryReader,
    StorageContext,
    load_index_from_storage,
    Settings,
    Document
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.llms.ollama import Ollama
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

# ChromaDB
import chromadb
from chromadb.config import Settings as ChromaSettings

# Утилиты
from tqdm import tqdm
from rich.console import Console
from rich.panel import Panel
from rich.table import Table

# Обработка документов
from pypdf import PdfReader
from docx import Document as DocxDocument

# Локальный конфиг
import config

console = Console()


class RAGEngine:
    """
    Класс для работы с RAG-системой

    Основные возможности:
    - Индексация документов из папки
    - Сохранение и загрузка индекса
    - Поиск по документам с помощью векторного хранилища
    - Генерация ответов с использованием LLM
    - Отслеживание источников информации
    """

    def __init__(self, verbose: bool = True):
        """
        Инициализация RAG Engine

        Args:
            verbose: Включить подробный вывод
        """
        self.verbose = verbose
        self.index = None
        self.query_engine = None
        self.last_sources = []
        self.document_count = 0

        # Создаем директорию для хранения, если её нет
        config.STORAGE_DIR.mkdir(exist_ok=True)

        self._log("Инициализация RAG Engine...", style="bold blue")

        # Инициализация компонентов
        self._initialize_llm()
        self._initialize_embeddings()
        self._initialize_vector_store()

        self._log("RAG Engine успешно инициализирован!", style="bold green")

    def _log(self, message: str, style: str = ""):
        """Вывод лога, если verbose=True"""
        if self.verbose:
            if style:
                console.print(f"[{style}]{message}[/{style}]")
            else:
                console.print(message)

    def _initialize_llm(self):
        """Инициализация LLM (Ollama с Llama 3.1)"""
        try:
            self._log("Подключение к Ollama LLM...", style="cyan")

            # Создаем LLM через Ollama
            self.llm = Ollama(
                model=config.OLLAMA_MODEL,
                base_url=config.OLLAMA_BASE_URL,
                request_timeout=config.OLLAMA_REQUEST_TIMEOUT,
                temperature=config.TEMPERATURE,
                system_prompt=config.SYSTEM_PROMPT
            )

            # Устанавливаем глобальные настройки LlamaIndex
            Settings.llm = self.llm

            # Проверяем доступность модели
            try:
                test_response = self.llm.complete("test")
                self._log(f"✓ Ollama модель {config.OLLAMA_MODEL} доступна", style="green")
            except Exception as e:
                raise ConnectionError(
                    f"Не удалось подключиться к Ollama. "
                    f"Убедитесь, что Ollama запущена и модель {config.OLLAMA_MODEL} загружена.\n"
                    f"Ошибка: {str(e)}"
                )

        except Exception as e:
            self._log(f"✗ Ошибка инициализации LLM: {str(e)}", style="red")
            raise

    def _initialize_embeddings(self):
        """Инициализация эмбеддинг модели"""
        try:
            self._log(f"Загрузка эмбеддинг модели {config.EMBEDDING_MODEL_NAME}...", style="cyan")

            # Загружаем русскоязычную эмбеддинг модель
            self.embed_model = HuggingFaceEmbedding(
                model_name=config.EMBEDDING_MODEL_NAME,
                device="cpu"  # Можно изменить на "cuda" если есть GPU
            )

            # Устанавливаем глобальные настройки
            Settings.embed_model = self.embed_model
            Settings.chunk_size = config.CHUNK_SIZE
            Settings.chunk_overlap = config.CHUNK_OVERLAP

            self._log(f"✓ Эмбеддинг модель загружена", style="green")

        except Exception as e:
            self._log(f"✗ Ошибка загрузки эмбеддинг модели: {str(e)}", style="red")
            raise

    def _initialize_vector_store(self):
        """Инициализация ChromaDB векторного хранилища"""
        try:
            self._log("Инициализация ChromaDB...", style="cyan")

            # Создаем директорию для ChromaDB
            os.makedirs(config.CHROMA_PERSIST_DIR, exist_ok=True)

            # Инициализируем ChromaDB клиент с персистентным хранилищем
            # Отключаем телеметрию для избежания предупреждений
            chroma_settings = ChromaSettings(
                anonymized_telemetry=False,
                allow_reset=True
            )

            chroma_client = chromadb.PersistentClient(
                path=config.CHROMA_PERSIST_DIR,
                settings=chroma_settings
            )

            # Получаем или создаем коллекцию
            try:
                # Пробуем удалить старую коллекцию при инициализации
                # (будет пересоздана при индексации)
                chroma_client.delete_collection(name=config.CHROMA_COLLECTION_NAME)
            except:
                pass

            self.chroma_collection = chroma_client.get_or_create_collection(
                name=config.CHROMA_COLLECTION_NAME
            )

            # Создаем ChromaVectorStore
            self.vector_store = ChromaVectorStore(
                chroma_collection=self.chroma_collection
            )

            self._log("✓ ChromaDB инициализирован", style="green")

        except Exception as e:
            self._log(f"✗ Ошибка инициализации ChromaDB: {str(e)}", style="red")
            raise

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Извлечение текста из PDF файла

        Args:
            file_path: Путь к PDF файлу

        Returns:
            Извлеченный текст
        """
        try:
            reader = PdfReader(str(file_path))
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    self._log(f"⚠ Ошибка чтения страницы {page_num} в {file_path.name}: {str(e)}", style="yellow")
                    continue

            return "\n\n".join(text_parts)

        except Exception as e:
            raise Exception(f"Ошибка извлечения текста из PDF: {str(e)}")

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        Извлечение текста из DOCX файла

        Args:
            file_path: Путь к DOCX файлу

        Returns:
            Извлеченный текст
        """
        try:
            doc = DocxDocument(str(file_path))
            text_parts = []

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            raise Exception(f"Ошибка извлечения текста из DOCX: {str(e)}")

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """
        Извлечение текста из TXT файла

        Args:
            file_path: Путь к TXT файлу

        Returns:
            Извлеченный текст
        """
        try:
            with open(file_path, 'r', encoding=config.ENCODING) as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Ошибка чтения TXT файла: {str(e)}")

    def _extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """
        Извлечение текста из файла в зависимости от его типа

        Args:
            file_path: Путь к файлу

        Returns:
            Извлеченный текст или None в случае ошибки
        """
        file_extension = file_path.suffix.lower()

        try:
            if file_extension == ".txt":
                return self._extract_text_from_txt(file_path)
            elif file_extension == ".pdf":
                return self._extract_text_from_pdf(file_path)
            elif file_extension == ".docx":
                return self._extract_text_from_docx(file_path)
            else:
                self._log(f"⚠ Неподдерживаемый формат файла: {file_path.name}", style="yellow")
                return None

        except Exception as e:
            self._log(f"⚠ Ошибка обработки {file_path.name}: {str(e)}", style="yellow")
            return None

    def index_documents(self, folder_path: str) -> bool:
        """
        Индексация документов из указанной папки

        Args:
            folder_path: Путь к папке с документами (txt, pdf, docx)

        Returns:
            True если индексация успешна, False иначе
        """
        try:
            folder_path = Path(folder_path)

            # Валидация пути
            if not folder_path.exists():
                self._log(f"✗ Папка не существует: {folder_path}", style="red")
                return False

            if not folder_path.is_dir():
                self._log(f"✗ Указанный путь не является папкой: {folder_path}", style="red")
                return False

            # Получаем список всех поддерживаемых файлов
            all_files = []
            for extension in config.SUPPORTED_FILE_EXTENSIONS:
                all_files.extend(folder_path.glob(f"*{extension}"))

            if not all_files:
                supported_formats = ", ".join(config.SUPPORTED_FILE_EXTENSIONS)
                self._log(f"✗ В папке нет поддерживаемых файлов ({supported_formats}): {folder_path}", style="red")
                return False

            # Показываем статистику по типам файлов
            file_types_count = {}
            for file_path in all_files:
                ext = file_path.suffix.lower()
                file_types_count[ext] = file_types_count.get(ext, 0) + 1

            stats_str = ", ".join([f"{ext}: {count}" for ext, count in file_types_count.items()])
            self._log(f"\nНайдено файлов: {len(all_files)} ({stats_str})", style="bold cyan")

            # Загружаем документы
            self._log("\nЗагрузка и обработка документов...", style="cyan")
            documents = []

            for file_path in tqdm(all_files, desc="Обработка файлов", disable=not config.SHOW_PROGRESS_BAR):
                try:
                    # Проверяем размер файла
                    if file_path.stat().st_size > config.MAX_FILE_SIZE:
                        self._log(f"⚠ Файл слишком большой, пропускаем: {file_path.name}", style="yellow")
                        continue

                    # Извлекаем текст в зависимости от типа файла
                    content = self._extract_text_from_file(file_path)

                    if content is None or not content.strip():
                        self._log(f"⚠ Не удалось извлечь текст из {file_path.name}", style="yellow")
                        continue

                    # Создаем документ с метаданными
                    doc = Document(
                        text=content,
                        metadata={
                            "file_name": file_path.name,
                            "file_path": str(file_path),
                            "file_size": file_path.stat().st_size,
                            "file_type": file_path.suffix.lower()
                        }
                    )
                    documents.append(doc)

                except Exception as e:
                    self._log(f"⚠ Ошибка обработки файла {file_path.name}: {str(e)}", style="yellow")
                    continue

            if not documents:
                self._log("✗ Не удалось загрузить ни одного документа", style="red")
                return False

            self.document_count = len(documents)
            self._log(f"✓ Загружено документов: {self.document_count}", style="green")

            # Создаем парсер для разбиения на чанки
            self._log("\nРазбиение документов на чанки...", style="cyan")
            text_splitter = SentenceSplitter(
                chunk_size=config.CHUNK_SIZE,
                chunk_overlap=config.CHUNK_OVERLAP
            )

            # Создаем индекс с ChromaDB
            self._log("Создание векторного индекса...", style="cyan")

            storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )

            self.index = VectorStoreIndex.from_documents(
                documents,
                storage_context=storage_context,
                transformations=[text_splitter],
                show_progress=config.SHOW_PROGRESS_BAR
            )

            # Сохраняем индекс
            self._log("Сохранение индекса...", style="cyan")
            self._save_index()

            # Создаем query engine
            self._create_query_engine()

            self._log("\n✓ Индексация завершена успешно!", style="bold green")
            return True

        except Exception as e:
            self._log(f"\n✗ Ошибка при индексации: {str(e)}", style="bold red")
            return False

    def _save_index(self):
        """Сохранение индекса на диск"""
        try:
            if self.index is None:
                return

            # Индекс ChromaDB автоматически персистентный
            # Дополнительно сохраняем метаданные
            metadata_path = config.STORAGE_DIR / "metadata.txt"
            with open(metadata_path, 'w') as f:
                f.write(f"documents_count={self.document_count}\n")

            self._log(f"✓ Индекс сохранен в: {config.STORAGE_DIR}", style="green")

        except Exception as e:
            self._log(f"⚠ Ошибка сохранения индекса: {str(e)}", style="yellow")

    def load_index(self) -> bool:
        """
        Загрузка существующего индекса

        Returns:
            True если индекс загружен успешно, False иначе
        """
        try:
            # Проверяем наличие сохраненного индекса
            if not config.STORAGE_DIR.exists():
                return False

            metadata_path = config.STORAGE_DIR / "metadata.txt"
            if not metadata_path.exists():
                return False

            self._log("Загрузка существующего индекса...", style="cyan")

            # Загружаем метаданные
            with open(metadata_path, 'r') as f:
                for line in f:
                    if line.startswith("documents_count="):
                        self.document_count = int(line.split("=")[1].strip())

            # Пересоздаем vector store с существующей коллекцией
            storage_context = StorageContext.from_defaults(
                vector_store=self.vector_store
            )

            # Создаем индекс из существующего storage
            self.index = VectorStoreIndex.from_vector_store(
                self.vector_store,
                storage_context=storage_context
            )

            # Создаем query engine
            self._create_query_engine()

            self._log(f"✓ Индекс загружен ({self.document_count} документов)", style="green")
            return True

        except Exception as e:
            self._log(f"⚠ Не удалось загрузить индекс: {str(e)}", style="yellow")
            return False

    def _create_query_engine(self):
        """Создание query engine для поиска"""
        if self.index is None:
            return

        self.query_engine = self.index.as_query_engine(
            similarity_top_k=config.SIMILARITY_TOP_K,
            response_mode="compact",
            verbose=False
        )

    def query(self, question: str) -> Dict[str, Any]:
        """
        Поиск ответа на вопрос

        Args:
            question: Вопрос пользователя

        Returns:
            Словарь с ответом и метаданными
        """
        if self.query_engine is None:
            return {
                "answer": "Индекс не загружен. Сначала проиндексируйте документы.",
                "sources": [],
                "error": True
            }

        try:
            # Выполняем запрос
            response = self.query_engine.query(question)

            # Извлекаем источники
            sources = []
            if hasattr(response, 'source_nodes'):
                for node in response.source_nodes:
                    source_info = {
                        "file_name": node.metadata.get("file_name", "Unknown"),
                        "score": node.score if hasattr(node, 'score') else 1.0,
                        "text_preview": node.text[:config.MAX_SOURCE_PREVIEW_LENGTH] + "..."
                                      if len(node.text) > config.MAX_SOURCE_PREVIEW_LENGTH
                                      else node.text
                    }
                    sources.append(source_info)

            self.last_sources = sources

            return {
                "answer": str(response),
                "sources": sources,
                "error": False
            }

        except Exception as e:
            return {
                "answer": f"Ошибка при обработке запроса: {str(e)}",
                "sources": [],
                "error": True
            }

    def get_sources(self) -> List[Dict[str, Any]]:
        """
        Получить источники последнего ответа

        Returns:
            Список источников
        """
        return self.last_sources

    def has_index(self) -> bool:
        """Проверка наличия загруженного индекса"""
        return self.index is not None

    def get_stats(self) -> Dict[str, Any]:
        """
        Получить статистику индекса

        Returns:
            Словарь со статистикой
        """
        return {
            "documents_count": self.document_count,
            "has_index": self.has_index(),
            "storage_path": str(config.STORAGE_DIR),
            "model": config.OLLAMA_MODEL,
            "embedding_model": config.EMBEDDING_MODEL_NAME
        }